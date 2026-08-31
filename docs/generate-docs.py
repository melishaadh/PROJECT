"""
Bundle DEVOPS_OVERVIEW.md + RUNBOOK.md into one TrekEasy-DevOps-Documentation
.docx and .pdf (both land next to this script in docs/).

Run:
    uv run --no-project --with markdown --with htmldocx --with python-docx \
        docs/generate-docs.py

Needs Google Chrome (or set CHROME below to msedge.exe) for the PDF step.
"""
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path

import markdown  # noqa
from htmldocx import HtmlToDocx  # noqa
from docx import Document  # noqa
from docx.shared import Pt, Inches  # noqa

OUT = Path(__file__).resolve().parent
CHROME = (
    os.environ.get("CHROME")
    or shutil.which("chrome")
    or r"C:\Program Files\Google\Chrome\Application\chrome.exe"
)

PARTS = [
    ("Part 1 — Project &amp; DevOps Overview", OUT / "DEVOPS_OVERVIEW.md"),
    ("Part 2 — DevOps Runbook", OUT / "RUNBOOK.md"),
]

CSS = """
@page { size: A4; margin: 22mm 18mm 20mm 18mm; }
* { box-sizing: border-box; }
body { font-family: 'Segoe UI', Calibri, Arial, sans-serif; font-size: 10.5pt;
       line-height: 1.5; color: #1a1a1a; }
h1 { font-size: 21pt; color: #0b5394; border-bottom: 3px solid #1b6ec2;
     padding-bottom: 4px; margin: 26pt 0 12pt; page-break-after: avoid; }
h2 { font-size: 15pt; color: #14477d; border-bottom: 1px solid #cfd8e3;
     padding-bottom: 3px; margin: 20pt 0 8pt; page-break-after: avoid; }
h3 { font-size: 12pt; color: #14477d; margin: 14pt 0 6pt; page-break-after: avoid; }
h4 { font-size: 10.5pt; color: #333; margin: 10pt 0 4pt; }
p, li { orphans: 3; widows: 3; }
code { font-family: 'Cascadia Mono', Consolas, monospace; font-size: 9pt;
       background: #f2f4f7; padding: 1px 4px; border-radius: 3px; }
pre { background: #f6f8fa; border: 1px solid #d7dde3; border-radius: 5px;
      padding: 9px 11px; overflow-x: auto; page-break-inside: avoid; }
pre code { background: none; padding: 0; font-size: 8.6pt; line-height: 1.45; }
table { border-collapse: collapse; width: 100%; margin: 8pt 0; font-size: 9pt;
        page-break-inside: avoid; }
th, td { border: 1px solid #c4ccd6; padding: 4px 7px; text-align: left;
         vertical-align: top; }
th { background: #eaf1f8; font-weight: 600; }
tr:nth-child(even) td { background: #f7f9fb; }
blockquote { border-left: 3px solid #1b6ec2; margin: 8pt 0; padding: 2pt 12pt;
             color: #444; background: #f4f8fc; }
a { color: #1b6ec2; text-decoration: none; }
hr { border: none; border-top: 1px solid #d0d7de; margin: 16pt 0; }
.title-page { text-align: center; padding-top: 32%; page-break-after: always; }
.title-page .t { font-size: 30pt; font-weight: 700; color: #14477d; }
.title-page .s { font-size: 13pt; color: #555; margin-top: 10pt; }
.title-page .m { font-size: 10pt; color: #888; margin-top: 40pt; }
h1 { color: #0b5394; }
"""

def md_to_html_fragment(path: Path) -> str:
    text = path.read_text(encoding="utf-8")
    # Drop the leading "skip to RUNBOOK" style cross-links that make no sense in a
    # single combined document, keep everything else.
    return markdown.markdown(
        text,
        extensions=["tables", "fenced_code", "toc", "sane_lists", "attr_list"],
    )

def build_html() -> str:
    body = [
        '<div class="title-page">',
        '<div class="t">TrekEasy</div>',
        '<div class="s">DevOps &amp; AWS ECS on Fargate — Complete Documentation</div>',
        '<div class="m">Project &amp; DevOps Overview &nbsp;•&nbsp; Manual Runbook</div>',
        "</div>",
    ]
    for heading, path in PARTS:
        body.append(f"<h1>{heading}</h1>")
        body.append(md_to_html_fragment(path))
        body.append('<hr style="page-break-after: always; border: none;">')
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        f"<style>{CSS}</style></head><body>{''.join(body)}</body></html>"
    )

def main() -> None:
    OUT.mkdir(exist_ok=True)
    html = build_html()
    html_path = OUT / "TrekEasy-DevOps-Documentation.html"
    html_path.write_text(html, encoding="utf-8")

    # ---- PDF via headless Chrome ----
    pdf_path = OUT / "TrekEasy-DevOps-Documentation.pdf"
    subprocess.run(
        [CHROME, "--headless", "--disable-gpu", "--no-pdf-header-footer",
         f"--print-to-pdf={pdf_path}", "--print-to-pdf-no-header",
         html_path.as_uri()],
        check=True, timeout=120,
    )
    print(f"PDF  -> {pdf_path}  ({pdf_path.stat().st_size // 1024} KB)")

    # ---- DOCX via htmldocx ----
    docx_path = OUT / "TrekEasy-DevOps-Documentation.docx"
    doc = Document()
    for s in doc.sections:
        s.page_width, s.page_height = Inches(8.27), Inches(11.69)  # A4
        s.left_margin = s.right_margin = Inches(0.8)
        s.top_margin = s.bottom_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    title = doc.add_heading("TrekEasy", level=0)
    doc.add_paragraph("DevOps & AWS ECS on Fargate — Complete Documentation")
    doc.add_paragraph("Part 1: Project & DevOps Overview   |   Part 2: Manual Runbook")
    doc.add_page_break()

    parser = HtmlToDocx()
    parser.table_style = "Table Grid"
    for heading, path in PARTS:
        doc.add_heading(re.sub("&amp;", "&", heading), level=1)
        parser.add_html_to_document(md_to_html_fragment(path), doc)
        doc.add_page_break()
    doc.save(docx_path)
    print(f"DOCX -> {docx_path}  ({docx_path.stat().st_size // 1024} KB)")

if __name__ == "__main__":
    sys.exit(main())
